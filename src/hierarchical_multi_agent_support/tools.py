"""
Tools for the multi-agent support system.
"""

import os
import logging
import requests
from typing import Dict, Any, Optional, List
from pathlib import Path
from abc import ABC, abstractmethod
from pydantic import BaseModel
import pypdf
from docx import Document
import openpyxl

from .config import Config
from .models import ToolResult
from .rag_search import RAGDocumentSearch


class BaseTool(ABC):
    """Base class for all tools."""

    def __init__(self, config: Config, logger: logging.Logger):
        """Initialize tool with configuration and logger."""
        self.config = config
        self.logger = logger

    @abstractmethod
    async def execute(self, **kwargs) -> ToolResult:
        """Execute the tool."""
        pass

    def _handle_error(self, error: Exception, operation: str) -> ToolResult:
        """Handle tool errors consistently."""
        error_msg = f"Error in {operation}: {str(error)}"
        self.logger.error(error_msg)
        return ToolResult(success=False, error=error_msg)


class WebSearchTool(BaseTool):
    """Tool for web search functionality."""

    def __init__(self, config: Config, logger: logging.Logger):
        """Initialize web search tool."""
        super().__init__(config, logger)
        self.timeout = config.tools.web_search.timeout
        self.max_results = config.tools.web_search.max_results
        self.enabled = config.tools.web_search.enabled

    async def execute(self, query: str) -> ToolResult:
        """Execute web search."""
        if not self.enabled:
            return ToolResult(
                success=False,
                error="Web search tool is disabled in configuration"
            )

        try:
            self.logger.info(f"Executing web search for query: {query}")

            # Validate input
            if not query or len(query.strip()) == 0:
                return ToolResult(
                    success=False,
                    error="Search query cannot be empty"
                )

            # Simulate web search (in a real implementation, you'd use a search API)
            # For demonstration, we'll return mock results
            search_results = self._simulate_web_search(query)

            self.logger.info(f"Web search completed successfully for query: {query}")
            return ToolResult(
                success=True,
                data=search_results,
                metadata={"query": query, "results_count": len(search_results)}
            )

        except requests.RequestException as e:
            return self._handle_error(e, "web search request")
        except Exception as e:
            return self._handle_error(e, "web search")

    def _simulate_web_search(self, query: str) -> List[Dict[str, str]]:
        """Simulate web search results."""
        # In a real implementation, this would use a search API like Google Custom Search
        mock_results = [
            {
                "title": f"Search result for '{query}' - Documentation",
                "url": f"https://docs.example.com/search?q={query}",
                "snippet": f"This is a mock search result for the query '{query}'. It provides relevant information about the topic."
            },
            {
                "title": f"Best practices for {query}",
                "url": f"https://bestpractices.example.com/{query}",
                "snippet": f"Learn about best practices and solutions related to {query}."
            },
            {
                "title": f"Common issues with {query}",
                "url": f"https://support.example.com/issues/{query}",
                "snippet": f"Troubleshooting guide for common issues related to {query}."
            }
        ]

        return mock_results[:self.max_results]


class RAGSearchTool(BaseTool):
    """Tool for RAG-based document search using semantic similarity."""

    def __init__(self, config: Config, logger: logging.Logger):
        """Initialize RAG search tool."""
        super().__init__(config, logger)
        self.enabled = config.tools.file_reader.enabled
        self.rag_search = RAGDocumentSearch(config, logger)

    async def execute(self, query: str, domain: str = "finance") -> ToolResult:
        """Execute semantic search through documents."""
        if not self.enabled:
            return ToolResult(
                success=False,
                error="RAG search tool is disabled in configuration"
            )

        try:
            self.logger.info(f"Executing RAG search for query: {query} in domain: {domain}")

            # Support both IT and Finance domains
            if domain not in ["finance", "it"]:
                return ToolResult(
                    success=False,
                    error="RAG search supports 'finance' and 'it' domains only"
                )

            # Get relevant context using RAG
            context = await self.rag_search.get_context_for_query(query, domain)

            if not context:
                return ToolResult(
                    success=False,
                    error="No relevant documents found for the query"
                )

            # Get the actual matching chunks for metadata
            relevant_chunks = await self.rag_search.search_documents(query, domain, top_k=5)

            metadata = {
                "query": query,
                "domain": domain,
                "chunks_found": len(relevant_chunks),
                "sources": list(set([chunk.source for chunk in relevant_chunks])),
                "similarity_scores": [chunk.metadata.get('similarity_score', 0.0) for chunk in relevant_chunks]
            }

            self.logger.info(f"RAG search found {len(relevant_chunks)} relevant chunks")
            return ToolResult(
                success=True,
                data=context,
                metadata=metadata
            )

        except Exception as e:
            return self._handle_error(e, "RAG search")


class ReadFileTool(BaseTool):
    """Tool for reading internal documentation files - deprecated in favor of RAG search."""

    def __init__(self, config: Config, logger: logging.Logger):
        """Initialize file reader tool."""
        super().__init__(config, logger)
        self.max_file_size = config.tools.file_reader.max_file_size
        self.allowed_extensions = config.tools.file_reader.allowed_extensions
        self.enabled = config.tools.file_reader.enabled
        self.it_docs_path = Path(config.documents.it_docs_path)
        self.finance_docs_path = Path(config.documents.finance_docs_path)

    async def execute(self, file_path: str, domain: str = "it") -> ToolResult:
        """Read a file from the internal documentation."""
        if not self.enabled:
            return ToolResult(
                success=False,
                error="File reader tool is disabled in configuration"
            )

        try:
            self.logger.info(f"Reading file: {file_path} from domain: {domain}")

            # Recommend using RAG search for both domains
            return ToolResult(
                success=False,
                error=f"Document reading now uses RAG search for better results. Please use the rag_search tool for {domain} queries."
            )

        except Exception as e:
            return self._handle_error(e, "file read")

    def _is_safe_path(self, file_path: Path, base_path: Path) -> bool:
        """Check if file path is safe (within base directory)."""
        try:
            file_path.resolve().relative_to(base_path.resolve())
            return True
        except ValueError:
            return False

    def _read_file_content(self, file_path: Path) -> str:
        """Read file content based on extension with support for multiple formats."""
        extension = file_path.suffix.lower()

        try:
            if extension in [".txt", ".md"]:
                return self._read_text_file(file_path)
            elif extension == ".pdf":
                return self._read_pdf_file(file_path)
            elif extension in [".doc", ".docx"]:
                return self._read_word_file(file_path)
            elif extension in [".xls", ".xlsx"]:
                return self._read_excel_file(file_path)
            else:
                # For unsupported file types, return a placeholder
                return f"[Content of {file_path.name} - Parser not implemented for {extension}]"
        except Exception as e:
            self.logger.error(f"Error reading file {file_path}: {str(e)}")
            return f"[Error reading {file_path.name}: {str(e)}]"

    def _read_text_file(self, file_path: Path) -> str:
        """Read plain text and markdown files."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def _read_pdf_file(self, file_path: Path) -> str:
        """Read PDF files and extract text content."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                text_content = []

                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num} ---\n{page_text}")
                    except Exception as e:
                        text_content.append(f"--- Page {page_num} ---\n[Error extracting text: {str(e)}]")

                if not text_content:
                    return "[PDF file appears to be empty or contains only images]"

                return "\n\n".join(text_content)
        except Exception as e:
            self.logger.error(f"Error reading PDF {file_path}: {str(e)}")
            return f"[Error reading PDF: {str(e)}]"

    def _read_word_file(self, file_path: Path) -> str:
        """Read Word documents (.doc, .docx)."""
        try:
            doc = Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_content.append("\n--- Table ---\n" + "\n".join(table_text))

            if not text_content:
                return "[Word document appears to be empty]"

            return "\n\n".join(text_content)
        except Exception as e:
            self.logger.error(f"Error reading Word document {file_path}: {str(e)}")
            return f"[Error reading Word document: {str(e)}]"

    def _read_excel_file(self, file_path: Path) -> str:
        """Read Excel files (.xls, .xlsx)."""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = [f"--- Sheet: {sheet_name} ---"]

                # Get all rows with data
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty rows
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):
                        sheet_content.append(" | ".join(row_data))

                if len(sheet_content) > 1:  # More than just the header
                    text_content.append("\n".join(sheet_content))

            if not text_content:
                return "[Excel file appears to be empty]"

            return "\n\n".join(text_content)
        except Exception as e:
            self.logger.error(f"Error reading Excel file {file_path}: {str(e)}")
            return f"[Error reading Excel file: {str(e)}]"


class ToolRegistry:
    """Registry for managing tools."""

    def __init__(self, config: Config, logger: logging.Logger):
        """Initialize tool registry."""
        self.config = config
        self.logger = logger
        self.tools: Dict[str, BaseTool] = {}
        self._register_tools()

    def _register_tools(self) -> None:
        """Register all available tools."""
        self.tools["web_search"] = WebSearchTool(self.config, self.logger)
        self.tools["read_file"] = ReadFileTool(self.config, self.logger)
        self.tools["rag_search"] = RAGSearchTool(self.config, self.logger)

        self.logger.info(f"Registered {len(self.tools)} tools")

    def get_tool(self, tool_name: str) -> Optional[BaseTool]:
        """Get a tool by name."""
        return self.tools.get(tool_name)

    def list_tools(self) -> List[str]:
        """List all available tools."""
        return list(self.tools.keys())

    async def execute_tool(self, tool_name: str, **kwargs) -> ToolResult:
        """Execute a tool by name."""
        tool = self.get_tool(tool_name)
        if not tool:
            return ToolResult(
                success=False,
                error=f"Tool not found: {tool_name}"
            )

        return await tool.execute(**kwargs)
